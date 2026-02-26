import argparse
import logging
import sys
from pathlib import Path

import pymupdf
from docx import Document
from ukrainian_word_stress import Stressifier, StressSymbol

SUFFIX = "_nagolos"
SUPPORTED_FORMATS = {".docx", ".pdf"}

# Configure logging
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)


def main() -> int:
    """Entry point for the command line interface.

    Parses CLI options, adjusts logging verbosity, and dispatches
    file processing. Intended for use in :mod:`__main__`.

    Returns:
        int: Exit code suitable for :func:`sys.exit` (0 on success,
            non-zero on failure).
    """
    args = parse_arguments()

    if args.verbose:
        logger.setLevel(logging.DEBUG)

    try:
        process_file(args.input_file, args.output)
        return 0
    except FileNotFoundError as e:
        logger.error("File error: %s", e)
        return 1
    except ValueError as e:
        logger.error("Input error: %s", e)
        return 1
    except Exception:
        logger.exception("Unexpected error")
        return 2


def process_file(input_file: str, output_file: str | None = None) -> None:
    """Read a DOCX or PDF file and write a stressed-output DOCX file.

    This is the core business logic. It validates the input path, ensures the
    file extension is supported, chooses a default output path when one is
    omitted, and delegates to :func:`process_docx` or :func:`process_pdf` based
    on the extension. The resulting :class:`docx.Document` is saved to the
    destination path.

    Parameters:
        input_file (str): Path to the input document. Supported suffixes are
            ``.docx`` and ``.pdf`` (case-insensitive).
        output_file (Optional[str]): Path to write the resulting DOCX file. If
            ``None`` the name is derived by appending ``_nagolos`` to the
            original stem.

    Returns:
        None: The function operates by side-effects and saves the output file.

    Raises:
        FileNotFoundError: When :paramref:`input_file` does not point to an
            existing regular file.
        ValueError: When the file has an unsupported extension.
        Exception: Propagated from lower-level processing routines such as
            file-read errors.
    """
    file_path = Path(input_file)

    # Validate input file exists
    if not file_path.is_file():
        raise FileNotFoundError(f"Input file not found: {input_file}")

    # Get file extension
    ext = file_path.suffix.lower()

    # Validate supported format
    if ext not in SUPPORTED_FORMATS:
        raise ValueError(
            f"Unsupported file format: {ext}. "
            f"Supported formats: {', '.join(SUPPORTED_FORMATS)}"
        )

    logger.info("Processing file: %s", file_path)

    # Determine output file path
    if output_file is None:
        output_file = file_path.with_name(f"{file_path.stem}{SUFFIX}.docx")

    out_doc = Document()
    stressify = Stressifier(stress_symbol=StressSymbol.CombiningAcuteAccent)

    if ext == ".docx":
        process_docx(file_path, out_doc, stressify)
    elif ext == ".pdf":
        process_pdf(file_path, out_doc, stressify)

    out_doc.save(output_file)
    logger.info("✓ Processing completed: %s", output_file)


def process_docx(file_path: Path, out_doc: Document, stressify: Stressifier) -> None:
    """Extract text from a DOCX and append stressed paragraphs.

    Each paragraph in ``file_path`` is stress-marked using the provided
    ``stressify`` callable, and the resulting text is added to ``out_doc``.

    Parameters:
        file_path (Path): Source DOCX path.
        out_doc (Document): The ``docx.Document`` object that will receive the
            transformed text. Mutated in place.
        stressify (Stressifier): Callable that accepts a string and returns the
            same string annotated with stress marks.

    Returns:
        None
    Raises:
        Exception: Propagated from file reading errors.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        logger.error("Error reading DOCX file: %s", e)
        raise
    for paragraph in doc.paragraphs:
        out_doc.add_paragraph(stressify(paragraph.text))


def process_pdf(file_path: Path, out_doc: Document, stressify: Stressifier) -> None:
    """Read text from a PDF and append stressed paragraphs to a DOCX.

    Iterates through all pages in ``file_path`` using ``pymupdf``; for each
    page the extracted text is passed through ``stressify`` before being added
    as a new paragraph to ``out_doc``.

    Parameters:
        file_path (Path): Source PDF path.
        out_doc (Document): The ``docx.Document`` object that will receive the
            transformed text. Mutated in place.
        stressify (Stressifier): Callable that accepts a string and returns the
            same string annotated with stress marks.

    Returns:
        None

    Raises:
        Exception: Propagated from file reading or text extraction errors.
    """
    try:
        with pymupdf.open(file_path) as doc:
            for page in doc:
                out_doc.add_paragraph(stressify(page.get_text()))
    except Exception as e:
        logger.error("Error reading PDF file: %s", e)
        raise


def parse_arguments() -> argparse.Namespace:
    """Configure and parse command-line options.

    Returns an :class:`argparse.Namespace` containing the following
    attributes:

    ``input_file`` (str)
        Path to the input document (required).
    ``output`` (Optional[str])
        Output path for the generated DOCX. Defaults to ``None`` which causes
        :func:`process_file` to pick a name.
    ``verbose`` (bool)
        When true, the root logger is set to DEBUG level before processing.

    Returns:
        argparse.Namespace: Parsed arguments ready for use by :func:`main`.
    """
    parser = argparse.ArgumentParser(
        description="Add stress marks to Ukrainian text in DOCX and PDF files.",
        epilog="Supported formats: DOCX, PDF. Output is always DOCX format.",
    )
    parser.add_argument("input_file", help="Input file (DOCX or PDF)")
    parser.add_argument(
        "-o",
        "--output",
        help="Output file path (default: input_file_nagolos.docx)",
        default=None,
    )
    parser.add_argument(
        "-v", "--verbose", action="store_true", help="Enable verbose logging"
    )

    return parser.parse_args()


if __name__ == "__main__":
    sys.exit(main())
