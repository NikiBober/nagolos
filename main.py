import os
from sys import argv

import pymupdf
from docx import Document
from ukrainian_word_stress import Stressifier, StressSymbol

filename = argv[1]
suffix = "_nagolos"
name, ext = os.path.splitext(filename)
new_filename = name + suffix + ".docx"

out = Document()
stressify = Stressifier(stress_symbol=StressSymbol.CombiningAcuteAccent)

if ext.lower() == ".docx":
    source = Document(filename)
    for paragraph in source.paragraphs:
        out.add_paragraph(stressify(paragraph.text))
elif ext.lower() == ".pdf":
    doc = pymupdf.open(filename)
    for page in doc:
        out.add_paragraph(stressify(page.get_text()))
    doc.close()

out.save(new_filename)
